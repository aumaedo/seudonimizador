import os
import hashlib
import pandas as pd
import streamlit as st
from docx import Document

def generate_pseudonym(name):
    return hashlib.sha256(name.encode()).hexdigest()[:10]

def process_docx(file_path, pseudonym_table):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for name in pseudonym_table.keys():
            if name in paragraph.text:
                paragraph.text = paragraph.text.replace(name, pseudonym_table[name])
    doc.save(file_path)

def process_files(files):
    pseudonym_table = {}
    
    for file in files:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            words = paragraph.text.split()
            for i in range(len(words) - 1):
                name = f"{words[i]} {words[i+1]}"
                if name not in pseudonym_table:
                    pseudonym_table[name] = generate_pseudonym(name)
    
    for file in files:
        process_docx(file.name, pseudonym_table)
    
    return pseudonym_table

def main():
    st.title("Seudonimizador de Informes Médicos")
    
    uploaded_files = st.file_uploader("Selecciona los archivos .docx", type="docx", accept_multiple_files=True)
    
    if uploaded_files:
        if st.button("Procesar archivos"):
            pseudonym_table = process_files(uploaded_files)
            
            st.success("Archivos procesados exitosamente")
            
            df = pd.DataFrame(list(pseudonym_table.items()), columns=["Nombre", "Seudónimo"])
            st.write("Tabla de conversión:")
            st.dataframe(df)
            
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button(
                "Descargar tabla de conversión",
                csv,
                "tabla_conversion.csv",
                "text/csv",
                key='download-csv'
            )

if __name__ == "__main__":
    main()
